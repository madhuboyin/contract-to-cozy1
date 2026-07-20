from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/product/ContractToCozy_Product_Framework.docx")

# Resolved design system: standard_business_brief with named brand overrides.
FONT = "Aptos"
INK = "203040"
NAVY = "17324D"
BLUE = "2E6F95"
TEAL = "2A7F78"
GOLD = "C7923E"
MUTED = "667785"
PALE_BLUE = "EAF2F7"
PALE_TEAL = "E8F4F2"
PALE_GOLD = "FBF3E4"
PALE_GRAY = "F3F5F7"
WHITE = "FFFFFF"
RED = "A34545"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_borders(table, color="D9E0E5", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def font_run(run, size=None, bold=None, color=INK, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_para(doc, text="", *, style=None, bold_lead=None, align=None, after=6, before=0, italic=False, color=INK):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        font_run(p.add_run(bold_lead), bold=True, color=color)
        font_run(p.add_run(text[len(bold_lead):]), italic=italic, color=color)
    else:
        font_run(p.add_run(text), italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    style = "C2C Bullet" if level == 0 else "C2C Bullet 2"
    p = doc.add_paragraph(style=style)
    font_run(p.add_run(text))
    return p


ACTIVE_NUMBER_ID = 203
NEXT_NUMBER_ID = 301


def add_number(doc, text, restart=False):
    global ACTIVE_NUMBER_ID, NEXT_NUMBER_ID
    if restart:
        numbering = doc.part.numbering_part.element
        n = OxmlElement("w:num")
        n.set(qn("w:numId"), str(NEXT_NUMBER_ID))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), "103")
        n.append(aid)
        lvl_override = OxmlElement("w:lvlOverride")
        lvl_override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        lvl_override.append(start_override)
        n.append(lvl_override)
        numbering.append(n)
        ACTIVE_NUMBER_ID = NEXT_NUMBER_ID
        NEXT_NUMBER_ID += 1
    p = doc.add_paragraph(style="C2C Number")
    ppr = p._p.get_or_add_pPr()
    existing = ppr.find(qn("w:numPr"))
    if existing is not None:
        ppr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(ACTIVE_NUMBER_ID))
    num_pr.extend([ilvl, nid])
    ppr.append(num_pr)
    font_run(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, color=fill, size="0")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    font_run(p.add_run(label.upper() + "  "), size=9.5, bold=True, color=accent)
    font_run(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, header_fill=NAVY, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        font_run(p.add_run(value), size=9.1, bold=True, color=WHITE)
    for r_idx, row in enumerate(rows):
        new_row = table.add_row()
        prevent_row_split(new_row)
        cells = new_row.cells
        for i, value in enumerate(row):
            cell = cells[i]
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFB")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            font_run(p.add_run(str(value)), size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_intro(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    font_run(p.add_run(text), size=11.5, color=MUTED)
    return p


def page_break(doc):
    doc.add_page_break()


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_keep_with_next(p)
    return p


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element

    def abstract(abstract_id, num_fmt, text, left, hanging, font=None):
        a = OxmlElement("w:abstractNum")
        a.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        a.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "100")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.extend([start, fmt, lvl_text, suff, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        a.append(lvl)
        numbering.append(a)

    def num(num_id, abstract_id):
        n = OxmlElement("w:num")
        n.set(qn("w:numId"), str(num_id))
        aid = OxmlElement("w:abstractNumId")
        aid.set(qn("w:val"), str(abstract_id))
        n.append(aid)
        numbering.append(n)

    abstract(101, "bullet", "•", 720, 360, "Arial")
    abstract(102, "bullet", "–", 1080, 360, "Arial")
    abstract(103, "decimal", "%1.", 720, 360)
    num(201, 101)
    num(202, 102)
    num(203, 103)

    styles = doc.styles
    for name, num_id in (("C2C Bullet", 201), ("C2C Bullet 2", 202), ("C2C Number", 203)):
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        ppr = style.element.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        nid = OxmlElement("w:numId")
        nid.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, nid])
        ppr.append(num_pr)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, MUTED, 0, 10),
        "Heading 1": (18, NAVY, 16, 8),
        "Heading 2": (14, BLUE, 12, 6),
        "Heading 3": (11.5, TEAL, 9, 4),
    }
    for name, (size, color, before, after) in specs.items():
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st.font.size = Pt(size)
        st.font.bold = name != "Subtitle"
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    configure_numbering(doc)


def set_headers_footers(doc):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        font_run(p.add_run("CONTRACTTOCOZY  /  PRODUCT FRAMEWORK"), size=8.5, bold=True, color=MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(0)
        font_run(fp.add_run("CONFIDENTIAL  •  STRATEGY & PRODUCT OPERATING MODEL  •  "), size=8, color=MUTED)
        add_page_field(fp)


def add_cover(doc):
    for _ in range(4):
        add_para(doc, "", after=12)
    p = add_para(doc, "PRODUCT FRAMEWORK", after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    font_run(p.runs[0], size=10, bold=True, color=GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    font_run(p.add_run("ContractToCozy"), size=32, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    font_run(p.add_run("From Home Records to Home Intelligence"), size=17, color=BLUE)
    add_callout(doc, "Product promise", "ContractToCozy turns the history and condition of a home into timely actions, confident decisions, and guided plans.", PALE_TEAL, TEAL)
    add_para(doc, "Stay ahead. Decide with confidence. Navigate major home moments.", align=WD_ALIGN_PARAGRAPH.CENTER, before=14, after=38, italic=True, color=MUTED)
    add_para(doc, "Comprehensive product strategy and operating framework", align=WD_ALIGN_PARAGRAPH.CENTER, after=4, color=MUTED)
    add_para(doc, "Version 1.0  |  July 2026", align=WD_ALIGN_PARAGRAPH.CENTER, after=0, color=MUTED)
    page_break(doc)


def add_contents(doc):
    heading(doc, "How to use this framework", 1)
    add_section_intro(doc, "This document aligns company narrative, product architecture, experience design, prioritization, measurement, and phased execution around one coherent homeowner value loop.")
    add_callout(doc, "Core decision", "Use the three customer jobs as the product operating model—not as three rigid navigation tabs. Organize the experience around the homeowner’s present situation and next best action.", PALE_GOLD, GOLD)
    heading(doc, "Framework map", 2)
    sections = [
        ("Part I", "Strategic foundation", "Vision, promise, positioning, customer wedge, and competitive thesis"),
        ("Part II", "Customer jobs", "Stay Ahead, Decide With Confidence, and Navigate Major Moments"),
        ("Part III", "Product system", "Living Home Record, intelligence, workflows, and action feedback loops"),
        ("Part IV", "Experience framework", "Onboarding, dashboard, navigation, recommendations, and journeys"),
        ("Part V", "Trust and governance", "Explainability, commercial integrity, privacy, safety, and quality"),
        ("Part VI", "Execution", "Prioritization, roadmap, metrics, operating model, risks, and decisions"),
        ("Appendices", "Reusable tools", "Feature test, recommendation contract, event template, scorecards, and glossary"),
    ]
    add_table(doc, ["Section", "Focus", "What it governs"], sections, [1300, 2500, 5560], font_size=9.5)
    heading(doc, "Decision status legend", 2)
    add_bullet(doc, "Committed principle — intended to remain stable across roadmap phases.")
    add_bullet(doc, "Recommended choice — the strongest current direction; validate with evidence.")
    add_bullet(doc, "Validation question — a material uncertainty that requires research or market proof.")
    page_break(doc)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_contents(doc)

    heading(doc, "Executive summary", 1)
    add_section_intro(doc, "ContractToCozy should evolve from a platform containing homeowner tools into a trusted home intelligence and action platform.")
    add_callout(doc, "Strategic thesis", "The defensible product is not a better filing cabinet or reminder app. It is a system that accumulates property context, recognizes what matters, helps the homeowner choose, guides completion, and learns from the result.", PALE_TEAL, TEAL)
    heading(doc, "The strategy in one page", 2)
    add_table(doc, ["Element", "Framework decision"], [
        ("Vision", "Every homeowner has a trusted guide that understands their home and helps protect it, improve it, and navigate important transitions."),
        ("Launch wedge", "Recommended: recent homebuyers in the first 24 months of ownership, activated through inspection and transaction information."),
        ("Recurring engine", "A personalized attention feed that identifies risks, deadlines, maintenance, savings, and unfinished work."),
        ("Differentiator", "Decision intelligence grounded in the home’s systems, history, documents, local context, and homeowner preferences."),
        ("Major-event strategy", "Build one reusable event engine; launch one or two journeys deeply rather than many shallow workflows."),
        ("Data asset", "A Living Home Record that becomes more complete and useful after every document, decision, repair, and project."),
        ("Trust model", "Recommendations expose facts, assumptions, rationale, confidence, commercial relationships, and professional-advice boundaries."),
        ("North-star outcome", "Important home actions are identified early and completed successfully."),
    ], [1900, 7460], font_size=9.6)
    heading(doc, "What this changes", 2)
    add_bullet(doc, "The dashboard begins with what needs attention—not a catalog of features.")
    add_bullet(doc, "Storage capabilities become inputs to action, recommendations, and continuity.")
    add_bullet(doc, "The three jobs govern strategy and prioritization, but do not become rigid silos.")
    add_bullet(doc, "The go-to-market wedge is narrow even though the platform vision remains broad.")
    add_bullet(doc, "Trust, explainability, and commercial integrity are designed into the recommendation system.")

    page_break(doc)
    heading(doc, "Part I — Strategic foundation", 1)
    add_section_intro(doc, "The strategic foundation defines the enduring purpose of ContractToCozy and the focused entry point through which the platform earns adoption.")
    heading(doc, "1. Vision, mission, and promise", 2)
    heading(doc, "Vision", 3)
    add_para(doc, "Every homeowner should have a trusted guide that understands their home and helps them protect it, improve it, and navigate every important transition.")
    heading(doc, "Mission", 3)
    add_para(doc, "Turn fragmented home information into timely, understandable, and trustworthy action throughout the life of the property.")
    heading(doc, "Product promise", 3)
    add_para(doc, "ContractToCozy turns the history and condition of a home into timely actions, confident decisions, and guided plans.")
    heading(doc, "Customer-facing narrative", 3)
    add_para(doc, "Stay ahead. Decide with confidence. Navigate major home moments.")
    heading(doc, "Long-form positioning", 3)
    add_para(doc, "ContractToCozy continuously understands the home, identifies what matters, helps the homeowner decide, guides what happens next, and learns from the outcome.")

    heading(doc, "2. The homeowner problem", 2)
    add_para(doc, "Homeownership is managed through disconnected documents, vendors, calendars, inboxes, memories, websites, and professional relationships. Important information exists, but it rarely arrives in the form of a clear next action.")
    add_bullet(doc, "Homeowners discover issues too late, after urgency removes options.")
    add_bullet(doc, "High-cost decisions are made without property-specific context.")
    add_bullet(doc, "Major events require coordination across tasks, documents, providers, deadlines, and dependencies.")
    add_bullet(doc, "The knowledge created by one repair or project is usually lost before the next decision.")
    add_bullet(doc, "Generic advice does not reflect the actual home, household constraints, or local environment.")
    add_callout(doc, "Opportunity", "Convert a fragmented, reactive experience into a continuous recognize–decide–act–learn loop.", PALE_BLUE, BLUE)

    heading(doc, "3. Category and differentiation", 2)
    add_para(doc, "ContractToCozy should define itself as a home intelligence and action platform. ‘Home management’ describes a category; it does not explain the distinctive value.")
    add_table(doc, ["Common category behavior", "ContractToCozy advantage"], [
        ("Store records", "Use records to detect deadlines, risk, eligibility, and decision context."),
        ("Send reminders", "Explain why an item matters, recommend the action, and guide completion."),
        ("Offer generic calculators", "Use property history, system age, costs, coverage, and preferences."),
        ("List service providers", "Connect provider choice to the decision, scope, quote, and recorded outcome."),
        ("Provide static checklists", "Adapt journeys to the property, event state, dependencies, and progress."),
    ], [3600, 5760])
    add_callout(doc, "Defensibility", "The moat is the combination of longitudinal property context, recommendation quality, workflow completion data, and earned homeowner trust—not any single utility.", PALE_GOLD, GOLD)

    heading(doc, "4. Initial customer and market wedge", 2)
    add_callout(doc, "Recommended launch wedge", "Recent homebuyers, especially first-time and early-tenure owners, within the first 24 months after purchase.", PALE_TEAL, TEAL)
    add_para(doc, "This segment encounters concentrated uncertainty, has a natural onboarding moment, possesses information-rich transaction artifacts, and immediately needs an ownership plan.")
    heading(doc, "Why this wedge is attractive", 3)
    add_bullet(doc, "High motivation: the homeowner is confronting unfamiliar systems, costs, and responsibilities.")
    add_bullet(doc, "Strong data bootstrap: inspection reports, disclosures, closing documents, warranties, and initial projects can seed the Living Home Record.")
    add_bullet(doc, "Natural conversion path: a buyer or move-in workflow can transition into ongoing maintenance and decision support.")
    add_bullet(doc, "High learning velocity: early ownership produces many observable decisions and completed actions.")
    add_bullet(doc, "Partner distribution potential: agents, inspectors, lenders, insurers, title/closing services, and employers.")
    heading(doc, "Primary persona", 3)
    add_table(doc, ["Dimension", "Working definition"], [
        ("Situation", "Purchased a primary residence within the last 24 months."),
        ("Mindset", "Wants to be responsible but lacks confidence, time, or a reliable operating system."),
        ("Pain", "Unclear priorities, scattered documents, unfamiliar systems, surprise costs, and uncertain providers."),
        ("Desired outcome", "A clear plan for what matters now, what can wait, what it may cost, and how to complete it."),
        ("Activation event", "Upload or connect the inspection report and confirm a small set of property facts."),
        ("First-value moment", "Receive a credible, personalized 90-day ownership plan within minutes."),
    ], [2200, 7160])
    heading(doc, "Validation criteria", 3)
    add_bullet(doc, "At least 60% of activated users complete the minimum property setup.")
    add_bullet(doc, "At least 50% identify one recommended action as new and useful.")
    add_bullet(doc, "At least 30% complete or intentionally dismiss one action in the first 30 days.")
    add_bullet(doc, "Partner acquisition produces lower-friction activation than cold direct-to-consumer acquisition.")
    add_para(doc, "These are initial hypothesis thresholds, not final targets; calibrate them after baseline research and pilot data.", italic=True, color=MUTED)

    heading(doc, "5. Expansion strategy", 2)
    add_para(doc, "The wedge determines where ContractToCozy starts, not where it ends. Expansion should follow adjacent needs that reuse the same property context and action engine.")
    add_table(doc, ["Stage", "Customer situation", "Expansion logic"], [
        ("1. Establish", "Recent buyer / new owner", "Create the property record, initial plan, and recurring attention loop."),
        ("2. Deepen", "Established homeowner", "Support maintenance, repair decisions, quotes, projects, renewals, and savings."),
        ("3. Orchestrate", "Owner entering a major event", "Guide renovation, damage recovery, sale preparation, or landlord transition."),
        ("4. Transfer", "Ownership or occupancy changes", "Preserve and selectively transfer the home’s history and active obligations."),
    ], [1400, 2600, 5360])

    page_break(doc)
    heading(doc, "Part II — Customer jobs", 1)
    add_section_intro(doc, "The three jobs describe why the homeowner opens ContractToCozy. They are benefit-oriented, non-exclusive, and connected through a shared property context.")
    heading(doc, "6. Job framework", 2)
    add_table(doc, ["Customer job", "Question", "Primary value", "Cadence"], [
        ("Stay Ahead", "What needs my attention now or soon?", "Fewer surprises and missed opportunities", "Ongoing / event-triggered"),
        ("Decide With Confidence", "What is the right thing to do?", "Clearer tradeoffs and better choices", "As needed / high intent"),
        ("Navigate Major Moments", "Something significant is happening—what do I do next?", "Coordinated progress from start to completion", "Episodic / multi-step"),
    ], [1900, 2650, 3000, 1810], font_size=9.1)
    add_callout(doc, "Operating rule", "Assign each capability one primary job for ownership and measurement, while permitting secondary contributions across the other jobs.", PALE_BLUE, BLUE)
    heading(doc, "The connected value loop", 2)
    add_para(doc, "UNDERSTAND THE HOME  →  RECOGNIZE WHAT MATTERS  →  HELP DECIDE  →  GUIDE ACTION  →  LEARN FROM THE OUTCOME", align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=10, color=TEAL)
    add_para(doc, "The three jobs are not merely destinations. Stay Ahead detects a need; Decide With Confidence resolves the choice; Navigate Major Moments coordinates execution when the work becomes a journey.")

    heading(doc, "7. Job 1 — Stay Ahead", 2)
    add_callout(doc, "Customer promise", "Know what needs attention before it becomes urgent.", PALE_TEAL, TEAL)
    add_para(doc, "Stay Ahead is the recurring engagement engine. It should create calm and control, not a stream of anxiety-producing alerts.")
    heading(doc, "Supported needs", 3)
    for x in ["Maintenance due or becoming due", "Renewals, expirations, warranties, and coverage deadlines", "Recalls, safety issues, property risks, and seasonal preparation", "Savings, rebates, tax incentives, and avoidable future costs", "Unfinished tasks, provider follow-ups, and unresolved recommendations", "Material changes to the home, local environment, or household plan"]:
        add_bullet(doc, x)
    heading(doc, "Attention-item contract", 3)
    add_table(doc, ["Required element", "Question answered"], [
        ("Signal", "What changed or became relevant?"),
        ("Why it matters", "What consequence or opportunity exists?"),
        ("Recommended action", "What should the homeowner do next?"),
        ("Timing", "When should it happen, and how urgent is it?"),
        ("Expected outcome", "What improves if the user acts?"),
        ("Confidence", "How certain is the system, and what is missing?"),
        ("Clear CTA", "Start, schedule, compare, upload, review, snooze, or dismiss."),
    ], [2400, 6960])
    heading(doc, "Guardrails", 3)
    add_bullet(doc, "Prioritize by consequence, urgency, confidence, household relevance, and actionability—not merely due date.")
    add_bullet(doc, "Bundle low-priority items into a weekly Home Brief; reserve immediate alerts for genuine urgency.")
    add_bullet(doc, "Permit ‘not relevant,’ ‘already done,’ and ‘remind me later’ feedback to improve future ranking.")
    add_bullet(doc, "Never imply that daily product use is necessary for responsible homeownership.")

    heading(doc, "8. Job 2 — Decide With Confidence", 2)
    add_callout(doc, "Customer promise", "Understand the options, tradeoffs, and best next action for this home and household.", PALE_BLUE, BLUE)
    add_para(doc, "Decision intelligence is the primary differentiation layer. It should synthesize property-specific context into a recommendation while preserving homeowner agency.")
    heading(doc, "Initial decision modules", 3)
    add_table(doc, ["Decision", "Inputs", "Output"], [
        ("Repair or replace", "Age, condition, repair history, warranty, efficiency, failure risk, costs", "Recommendation, timing, break-even logic, uncertainty"),
        ("Evaluate a quote", "Scope, line items, local ranges, system context, exclusions, provider history", "Reasonableness, gaps, questions to ask, comparison"),
        ("Prioritize a project", "Risk, urgency, dependencies, budget, enjoyment, value impact", "Ranked plan and sequencing"),
        ("Act now or wait", "Risk trajectory, seasonality, availability, price trend, household constraints", "Timing recommendation and trigger conditions"),
        ("Choose a provider", "Fit, credentials, scope, availability, history, reviews, commercial disclosures", "Shortlist and selection rationale"),
    ], [1800, 4200, 3360], font_size=8.8)
    heading(doc, "Decision experience stages", 3)
    add_number(doc, "Frame the decision in the homeowner’s language.", restart=True)
    add_number(doc, "Confirm the relevant facts and expose missing information.")
    add_number(doc, "Present realistic options, including doing nothing or waiting when appropriate.")
    add_number(doc, "Compare cost, risk, timing, effort, coverage, and household preferences.")
    add_number(doc, "Recommend an option with rationale and confidence—not false precision.")
    add_number(doc, "Convert the choice into an action plan, provider workflow, or future trigger.")
    add_number(doc, "Record the decision and result in the Living Home Record.")

    heading(doc, "9. Job 3 — Navigate Major Moments", 2)
    add_callout(doc, "Customer promise", "Follow a clear, adaptive path from start to completion when something significant happens.", PALE_GOLD, GOLD)
    add_para(doc, "A major moment is not one decision. It is a stateful journey involving tasks, documents, milestones, dependencies, decisions, providers, communications, and changing circumstances.")
    heading(doc, "Candidate journeys", 3)
    add_bullet(doc, "Buying and moving in")
    add_bullet(doc, "Major renovation")
    add_bullet(doc, "Storm damage and insurance claim")
    add_bullet(doc, "Emergency repair and recovery")
    add_bullet(doc, "Preparing to sell and selling")
    add_bullet(doc, "Renting the property or becoming a landlord")
    add_bullet(doc, "Estate, ownership, or occupancy transfer")
    heading(doc, "Journey contract", 3)
    add_table(doc, ["Component", "Role"], [
        ("Outcome definition", "States what ‘complete’ means for the homeowner."),
        ("Stages and milestones", "Create an understandable progress model."),
        ("Adaptive next step", "Surfaces the highest-value action based on current state."),
        ("Dependencies", "Prevent premature steps and expose blockers."),
        ("Decision points", "Invoke decision intelligence at moments of choice."),
        ("Documents and evidence", "Collect what is required and preserve the event record."),
        ("People and providers", "Clarify responsibility, communication, and follow-up."),
        ("Completion and handoff", "Close the event and feed durable facts into ongoing ownership."),
    ], [2450, 6910])
    heading(doc, "Initial event selection", 3)
    add_para(doc, "Recommended first journey: Buying and Moving In, because it aligns directly with the launch wedge and creates the initial property intelligence foundation. Recommended second journey for validation: Preparing to Sell, because it reuses property history and has clear milestones and economic value.")
    add_callout(doc, "Scope discipline", "Build the reusable event engine once, but do not launch every candidate event. A journey should be deep enough to create confidence and completion, not merely provide a checklist.", PALE_GOLD, GOLD)

    page_break(doc)
    heading(doc, "Part III — Product system", 1)
    add_section_intro(doc, "The platform capabilities sit behind the customer jobs. Their purpose is to improve recognition, recommendations, completion, and learning over time.")
    heading(doc, "10. Three-layer product architecture", 2)
    add_table(doc, ["Layer", "Purpose", "Examples"], [
        ("1. Customer jobs", "Why the homeowner engages", "Stay Ahead; Decide With Confidence; Navigate Major Moments"),
        ("2. Customer outcomes", "What successful use produces", "Avoid a missed deadline; choose a repair path; complete a claim stage"),
        ("3. Platform capabilities", "How the product creates the outcome", "Property profile, documents, rules, AI, workflow, providers, notifications"),
    ], [1900, 3000, 4460])
    add_para(doc, "Users should understand the first two layers. The third layer should remain mostly invisible until a capability becomes relevant to the task at hand.")

    heading(doc, "11. The Living Home Record", 2)
    add_callout(doc, "Strategic asset", "A continuously improving, permissioned record of the home’s systems, documents, events, decisions, work, costs, providers, and changing condition.", PALE_TEAL, TEAL)
    add_para(doc, "The Living Home Record is more than storage. Each fact should have provenance, recency, confidence, ownership, and permitted uses so it can safely support future recommendations.")
    heading(doc, "Core domains", 3)
    add_table(doc, ["Domain", "Representative information", "Product value"], [
        ("Property", "Address, type, age, size, location, occupancy", "Baseline context and eligibility"),
        ("Systems & components", "HVAC, roof, plumbing, electrical, appliances", "Maintenance, risk, lifecycle, repair decisions"),
        ("Documents", "Inspection, invoice, warranty, policy, permit, quote", "Fact extraction, evidence, deadlines, continuity"),
        ("Timeline", "Purchases, work, incidents, changes, decisions", "Longitudinal context and transferability"),
        ("Financial", "Costs, budgets, coverage, rebates, value estimates", "Affordability, tradeoffs, savings, planning"),
        ("Providers", "Contacts, scope, quotes, credentials, outcomes", "Selection, follow-up, repeat work"),
        ("Household preferences", "Budget posture, risk tolerance, priorities, timing", "Personalized recommendations"),
        ("Local context", "Climate, regulations, incentives, risks, market costs", "Relevance and localization"),
    ], [1750, 3600, 4010], font_size=8.8)
    heading(doc, "Data quality model", 3)
    add_bullet(doc, "Source: user entry, document extraction, provider input, public data, partner data, or system inference.")
    add_bullet(doc, "Confidence: verified, likely, inferred, disputed, or unknown.")
    add_bullet(doc, "Freshness: when the fact was observed and when it should be reviewed.")
    add_bullet(doc, "Lineage: which document, person, integration, or calculation produced it.")
    add_bullet(doc, "Permission: which people, features, partners, or future owners may use or receive it.")
    add_bullet(doc, "Supersession: how newer facts replace or coexist with historical facts.")

    heading(doc, "12. Intelligence and orchestration layers", 2)
    add_table(doc, ["Capability", "Responsibility"], [
        ("Signal engine", "Detect due dates, change, risk, opportunity, eligibility, and incomplete work."),
        ("Prioritization engine", "Rank attention items using consequence, urgency, relevance, confidence, and effort."),
        ("Decision engine", "Generate options, compare tradeoffs, recommend, explain, and record the choice."),
        ("Workflow engine", "Manage stages, dependencies, tasks, ownership, documents, and completion."),
        ("Personalization", "Apply household goals, constraints, preferences, and interaction feedback."),
        ("Notification system", "Deliver the right message through the right channel and cadence."),
        ("Provider layer", "Support discovery, qualification, quotes, coordination, completion, and outcome history."),
        ("AI interaction layer", "Extract, summarize, answer, explain, and assist while respecting evidence boundaries."),
    ], [2500, 6860])
    heading(doc, "Role of AI", 3)
    add_para(doc, "AI is an enabling layer, not the customer proposition. It should reduce effort, make context understandable, and help compose actions while preserving deterministic controls for deadlines, permissions, money, safety, and regulated advice.")
    add_bullet(doc, "Good uses: document extraction, summarization, question answering, option framing, explanation, plan drafting, and conversational navigation.")
    add_bullet(doc, "Controlled uses: risk scoring, financial recommendations, claim guidance, provider ranking, and any recommendation with material consequences.")
    add_bullet(doc, "Poor use: an ungrounded general assistant presented as if it understands the property without evidence.")

    heading(doc, "13. Closed-loop learning", 2)
    add_para(doc, "Every meaningful interaction should create useful state, not merely a transient screen view.")
    add_table(doc, ["User action", "State captured", "Future benefit"], [
        ("Completes maintenance", "Date, provider, cost, notes, evidence", "Next interval, system history, future quote context"),
        ("Dismisses an alert", "Reason and relevance feedback", "Better prioritization and fewer noisy alerts"),
        ("Chooses repair", "Option, rationale, scope, expected outcome", "Lifecycle knowledge and later replace decision"),
        ("Uploads a quote", "Provider, line items, exclusions, pricing", "Comparison, market learning, project record"),
        ("Completes an event stage", "Milestone, documents, decisions, responsible parties", "Adaptive next step and transferable history"),
    ], [2100, 3000, 4260], font_size=9.0)

    page_break(doc)
    heading(doc, "Part IV — Experience framework", 1)
    add_section_intro(doc, "The experience should feel situation-first and action-oriented. The three jobs govern design without forcing the homeowner to understand the product taxonomy.")
    heading(doc, "14. Experience principles", 2)
    principles = [
        ("Lead with what matters", "Prioritize the next useful action over exhaustive information."),
        ("Earn context progressively", "Deliver value with minimal setup; request more only when the benefit is clear."),
        ("Explain, then recommend", "Show the property facts and reasoning that support an action."),
        ("Preserve agency", "Offer options, editability, deferral, dismissal, and escalation to professionals."),
        ("Coordinate continuity", "Carry facts and progress across reminders, decisions, providers, and events."),
        ("Reduce anxiety", "Use calm language, meaningful urgency, and clear resolution paths."),
        ("Finish the loop", "Every insight should lead to action, intentional deferral, or a recorded outcome."),
        ("Design for absence", "The product remains valuable even when the homeowner does not open it daily."),
    ]
    for title, body in principles:
        add_para(doc, f"{title}. {body}", bold_lead=f"{title}.")

    heading(doc, "15. Cold-start and onboarding", 2)
    add_callout(doc, "Onboarding principle", "Deliver useful guidance from minimal information, then earn additional context progressively.", PALE_BLUE, BLUE)
    heading(doc, "Recommended onboarding flow", 3)
    add_number(doc, "Establish intent: confirm whether the user just bought, already owns, is planning a project, or is responding to an event.", restart=True)
    add_number(doc, "Identify the property: address and occupancy establish an initial baseline.")
    add_number(doc, "Import the highest-value artifact: inspection report for the launch wedge; otherwise use a relevant document or short guided setup.")
    add_number(doc, "Confirm only the facts required to avoid misleading guidance.")
    add_number(doc, "Generate an immediate personalized output: the first 90-day ownership plan and top three attention items.")
    add_number(doc, "Invite progressive enrichment tied to concrete benefits, such as warranty alerts or quote evaluation.")
    heading(doc, "Progressive context ladder", 3)
    add_table(doc, ["Level", "User effort", "Value unlocked"], [
        ("0. Address", "Seconds", "Property baseline, local context, generic seasonal relevance"),
        ("1. Key facts", "1–3 minutes", "Basic prioritization and system lifecycle guidance"),
        ("2. One document", "Upload / connect", "Extracted risks, dates, systems, and personalized actions"),
        ("3. Outcome history", "Created through use", "Better recommendations, cost context, and continuity"),
        ("4. Connected ecosystem", "Explicit permission", "Automated updates, provider coordination, and lower manual effort"),
    ], [1300, 2200, 5860])
    heading(doc, "First-value standard", 3)
    add_para(doc, "Within one session, a new user should receive something credible and actionable that could not be produced from a generic home-maintenance checklist alone.")

    heading(doc, "16. Dashboard and navigation", 2)
    heading(doc, "Default dashboard hierarchy", 3)
    add_number(doc, "What needs your attention — ranked items with rationale, timing, outcome, confidence, and CTA.", restart=True)
    add_number(doc, "Decisions you may need to make — contextual entry points based on property state and active work.")
    add_number(doc, "Active major moment — shown prominently only when an event is in progress.")
    add_number(doc, "Home at a glance — health, recent activity, open work, and coverage gaps as supporting context.")
    add_number(doc, "Ask about this home — conversational access grounded in the Living Home Record.")
    add_callout(doc, "Important", "Do not duplicate the jobs as three top-level tabs while also retaining a full feature-led navigation. That creates two competing information architectures.", PALE_GOLD, GOLD)
    heading(doc, "Recommended navigation model", 3)
    add_table(doc, ["Surface", "Role"], [
        ("Home", "Attention, recommended next action, active work, and brief status."),
        ("Plan & Projects", "Decisions, tasks, projects, budgets, and active event journeys."),
        ("Home Record", "Systems, timeline, documents, work history, providers, and property facts."),
        ("Ask", "Situation-first conversational entry grounded in property context."),
        ("Search / command", "Direct access to objects and actions without understanding taxonomy."),
    ], [2000, 7360])
    add_para(doc, "The exact labels should be usability-tested. The governing principle is stable: visible navigation should reflect homeowner intent and continuity, not internal services.", italic=True, color=MUTED)

    heading(doc, "17. Recommendation experience", 2)
    add_para(doc, "A recommendation is a product object with evidence, reasoning, action, state, and feedback—not a paragraph generated on demand.")
    add_table(doc, ["Field", "Requirement"], [
        ("Recommendation", "A concise proposed action or choice."),
        ("Outcome", "What the recommendation is intended to protect, improve, save, or complete."),
        ("Evidence", "Relevant verified facts and their sources."),
        ("Assumptions", "Estimates or inferred facts that materially affect the recommendation."),
        ("Options", "Realistic alternatives, including wait or do nothing when valid."),
        ("Tradeoffs", "Cost, risk, timing, effort, coverage, dependencies, and household preferences."),
        ("Confidence", "Clear level with an explanation of what would increase certainty."),
        ("Commercial disclosure", "Any referral, sponsorship, ranking, or revenue relationship."),
        ("Next action", "A specific CTA and expected subsequent state."),
        ("Feedback", "Accept, modify, defer, dismiss, correct, or request professional review."),
    ], [2250, 7110], font_size=9.0)
    heading(doc, "Urgency language", 3)
    add_table(doc, ["Level", "Meaning", "Experience"], [
        ("Now", "Likely safety, active damage, legal deadline, or rapidly compounding loss", "Immediate alert and clear escalation"),
        ("Soon", "Time-sensitive within a defined window", "Prominent attention item with due date"),
        ("Plan", "Expected future cost, lifecycle event, or seasonal task", "Budget/timeline planning and reminder"),
        ("Consider", "Optional opportunity or preference-dependent improvement", "Low-pressure suggestion with rationale"),
    ], [1200, 4850, 3310], font_size=9.0)

    heading(doc, "18. Notification and cadence model", 2)
    add_para(doc, "Notifications should create timely action without training users to ignore the product.")
    add_bullet(doc, "Immediate: active damage, safety, recall, expiring rights, or high-confidence urgent deadline.")
    add_bullet(doc, "Weekly Home Brief: grouped due-soon tasks, progress, savings, and unresolved decisions.")
    add_bullet(doc, "Monthly outlook: upcoming maintenance, lifecycle planning, budget needs, and changing risks.")
    add_bullet(doc, "Event-triggered: adaptive next step, dependency resolved, provider response, or stage deadline.")
    add_bullet(doc, "User-controlled: channel, quiet hours, digest frequency, topics, and escalation preferences.")

    page_break(doc)
    heading(doc, "Part V — Trust and governance", 1)
    add_section_intro(doc, "ContractToCozy earns the right to guide consequential decisions only by being transparent about evidence, uncertainty, incentives, permissions, and professional boundaries.")
    heading(doc, "19. Trust framework", 2)
    add_table(doc, ["Trust dimension", "Product requirement"], [
        ("Accuracy", "Use verified facts where possible; identify stale, inferred, or conflicting data."),
        ("Explainability", "Show why the recommendation applies to this home and household."),
        ("Uncertainty", "Express confidence honestly and identify missing information."),
        ("Agency", "Let users inspect, correct, defer, dismiss, or choose another option."),
        ("Professional boundaries", "Distinguish education and decision support from licensed or emergency advice."),
        ("Commercial integrity", "Disclose incentives and keep analysis independent from monetization."),
        ("Privacy", "Make collection, use, sharing, retention, and transfer understandable and controllable."),
        ("Accountability", "Record material recommendations, evidence, versions, user decisions, and outcomes."),
    ], [2350, 7010])

    heading(doc, "20. Recommendation safety tiers", 2)
    add_table(doc, ["Tier", "Examples", "Controls"], [
        ("Low consequence", "Organization, routine reminders, document summaries", "Source links, correction, basic confidence"),
        ("Material financial", "Repair/replace, project prioritization, quote evaluation", "Evidence display, ranges, assumptions, audit log, human escalation"),
        ("Regulated / coverage", "Insurance, refinancing, claims, legal or tax implications", "Educational framing, jurisdiction checks, licensed-professional boundary, enhanced review"),
        ("Safety / emergency", "Gas smell, electrical hazard, active flooding, structural risk", "Conservative escalation, emergency guidance, no delay-causing workflow"),
    ], [1500, 3300, 4560], font_size=8.9)
    add_callout(doc, "Non-negotiable", "ContractToCozy must never allow conversational fluency to appear more certain than the underlying evidence.", "F9ECEC", RED)

    heading(doc, "21. Commercial integrity", 2)
    add_para(doc, "Provider referrals, financing, insurance, sponsorships, or commerce may support the business model, but they can undermine the central promise if incentives are hidden or influence the analysis.")
    add_bullet(doc, "Recommendations prioritize the homeowner’s stated interests and the evidence available.")
    add_bullet(doc, "Sponsored placement is visually and textually distinct from analytical ranking.")
    add_bullet(doc, "Compensation, referral, or ownership relationships are disclosed before the user acts.")
    add_bullet(doc, "Users can view non-commercial alternatives and understand selection criteria.")
    add_bullet(doc, "Marketplace conversion does not define recommendation quality; successful homeowner outcomes do.")
    add_bullet(doc, "A governance owner reviews conflicts, complaints, ranking behavior, and outcome disparities.")

    heading(doc, "22. Privacy, permission, and transfer", 2)
    add_para(doc, "A home record can contain sensitive financial, occupancy, security, insurance, and behavioral information. Permission design must reflect the difference between the property’s history and the current household’s private data.")
    add_table(doc, ["Principle", "Implementation direction"], [
        ("Purpose limitation", "Collect and use information for clear homeowner value, not speculative future use."),
        ("Progressive consent", "Request permissions at the moment a connected benefit is understandable."),
        ("Household roles", "Support owner, co-owner, occupant, professional, and temporary collaborator access."),
        ("Selective sharing", "Share only the necessary records with a provider, partner, or professional."),
        ("Transfer boundaries", "Separate transferable property history from non-transferable household data."),
        ("Deletion and correction", "Provide understandable controls and preserve required audit obligations."),
        ("Security by sensitivity", "Apply stronger safeguards to access, financial, insurance, and occupancy data."),
    ], [2300, 7060])

    page_break(doc)
    heading(doc, "Part VI — Execution framework", 1)
    add_section_intro(doc, "Execution translates the vision into explicit sequencing, measurable outcomes, and governance decisions.")
    heading(doc, "23. Product prioritization model", 2)
    add_para(doc, "Features should be prioritized by the homeowner outcome and compounding platform value they create—not by the number of surfaces they add.")
    heading(doc, "Weighted score", 3)
    add_table(doc, ["Criterion", "Weight", "Evaluation question"], [
        ("Outcome value", "20%", "Does it prevent loss, reduce cost, increase confidence, or complete meaningful work?"),
        ("Frequency or stakes", "15%", "Is the need frequent, urgent, expensive, or emotionally significant?"),
        ("Context advantage", "15%", "Is ContractToCozy materially better because it knows this home?"),
        ("Actionability", "15%", "Does it produce a clear next action and measurable outcome?"),
        ("Learning value", "10%", "Does use improve the Living Home Record or future recommendations?"),
        ("Retention / continuity", "10%", "Does it create a reason to return or sustain progress?"),
        ("Strategic reuse", "10%", "Does it strengthen multiple jobs or reusable platform capabilities?"),
        ("Trust / feasibility", "5%", "Can it be delivered credibly, safely, and with available data?"),
    ], [2100, 1000, 6260], font_size=9.0)
    add_para(doc, "Score each criterion from 1 to 5, multiply by its weight, and require a named owner to document evidence and assumptions. Use effort, dependency, and risk as sequencing constraints after value scoring rather than allowing ease alone to dominate.")
    heading(doc, "Automatic downgrade signals", 3)
    add_bullet(doc, "Primarily displays stored information without helping the homeowner act.")
    add_bullet(doc, "Duplicates a generic utility without a property-context advantage.")
    add_bullet(doc, "Has no measurable homeowner outcome or feedback loop.")
    add_bullet(doc, "Creates a new destination without strengthening the core value loop.")
    add_bullet(doc, "Requires high setup effort before demonstrating value.")
    add_bullet(doc, "Introduces material trust risk without sufficient evidence or governance controls.")

    heading(doc, "24. Phased product roadmap", 2)
    add_callout(doc, "Roadmap principle", "Sequence capabilities around a complete value loop for one customer segment before expanding platform breadth.", PALE_TEAL, TEAL)
    add_table(doc, ["Phase", "Objective", "Core deliverables", "Exit evidence"], [
        ("0. Validate", "Prove problem, wedge, and first-value proposition", "Research; concierge pilot; inspection ingestion prototype; recommendation tests", "Users trust and act on personalized priorities"),
        ("1. Establish", "Create the Living Home Record and first ownership plan", "Address baseline; inspection import; fact confirmation; 90-day plan; attention feed", "Activation and first-action completion meet thresholds"),
        ("2. Engage", "Build the recurring Stay Ahead loop", "Maintenance; deadlines; weekly brief; feedback; home timeline; notifications", "Sustained useful return without alert fatigue"),
        ("3. Differentiate", "Deliver high-quality decision intelligence", "Repair/replace; quote evaluation; project priority; evidence and confidence model", "Recommendations improve confidence and lead to recorded outcomes"),
        ("4. Orchestrate", "Guide one major event deeply", "Buying/move-in journey; event engine; documents; tasks; dependencies; providers", "Meaningful milestone completion and event-to-recurring conversion"),
        ("5. Expand", "Add adjacent segments and journeys", "Seller Prep validation; partner distribution; selective marketplace and transfer", "Reuse lowers cost and preserves recommendation trust"),
    ], [1250, 2200, 3650, 2260], font_size=8.2)
    heading(doc, "Not-now list", 3)
    add_bullet(doc, "A broad provider marketplace before ContractToCozy can frame scope and evaluate outcomes.")
    add_bullet(doc, "Many shallow life-event checklists marketed as complete guidance.")
    add_bullet(doc, "Daily-engagement mechanics that create noise without homeowner value.")
    add_bullet(doc, "A general-purpose AI assistant disconnected from verified home context.")
    add_bullet(doc, "Data collection that cannot be tied to a visible current or near-term benefit.")

    heading(doc, "25. Measurement framework", 2)
    add_callout(doc, "North-star outcome", "The percentage of important home actions identified early and completed successfully.", PALE_BLUE, BLUE)
    add_para(doc, "Because ‘important,’ ‘early,’ and ‘successfully’ require operational definitions, the north star should be implemented as a transparent composite supported by job-level metrics.")
    add_table(doc, ["Metric family", "Example measures"], [
        ("Activation", "Property identified; first high-value artifact imported; critical facts confirmed; first plan viewed"),
        ("Stay Ahead", "High-priority actions surfaced early; completion before due date; alert usefulness; noise/dismissal rate"),
        ("Decide", "Decision confidence improvement; recommendation usefulness; option selected; outcome recorded; override reasons"),
        ("Major Moments", "Milestones completed on time; blocked-time reduction; journey completion; transition into ongoing ownership"),
        ("Living Home Record", "Verified fact growth; document extraction confirmation; freshness; correction rate; reusable context coverage"),
        ("Trust", "Explanation comprehension; confidence calibration; complaints; recommendation reversals; disclosure understanding"),
        ("Business", "Qualified activation cost; partner conversion; retention; paid conversion; revenue per successful outcome"),
    ], [2200, 7160], font_size=8.9)
    heading(doc, "Metrics to avoid optimizing in isolation", 3)
    add_bullet(doc, "Daily active users: responsible homeownership does not require daily app use.")
    add_bullet(doc, "Notifications sent: volume can increase while trust and usefulness decline.")
    add_bullet(doc, "Documents stored: inventory is not an outcome unless it improves action or continuity.")
    add_bullet(doc, "Provider leads: monetization can conflict with homeowner value if disconnected from successful outcomes.")
    add_bullet(doc, "AI messages: conversational activity is not proof of correct or useful guidance.")

    heading(doc, "26. Product operating model", 2)
    add_para(doc, "The three jobs should shape planning, ownership, discovery, design review, launch readiness, and measurement.")
    add_table(doc, ["Operating practice", "Required behavior"], [
        ("Quarterly planning", "Allocate investment across the end-to-end value loop and declare the primary job/outcome for each initiative."),
        ("Feature brief", "State job, outcome, user action, system learning, evidence, trust tier, and success metric."),
        ("Design review", "Test situation-first comprehension, action clarity, explainability, correction, and resolution path."),
        ("Recommendation review", "Include product, domain, data, trust/safety, legal/compliance, and commercial-integrity perspectives as needed."),
        ("Launch gate", "Verify data quality, fallback behavior, observability, support readiness, and outcome measurement."),
        ("Post-launch learning", "Review completed outcomes, overrides, dismissals, errors, complaints, and segment differences."),
    ], [2300, 7060], font_size=9.0)
    heading(doc, "Recommended ownership model", 3)
    add_bullet(doc, "Experience pillar owners: accountable for a customer job and its outcomes, not a screen collection.")
    add_bullet(doc, "Platform capability owners: accountable for reusable records, intelligence, workflows, notifications, identity, and permissions.")
    add_bullet(doc, "Trust owner or council: accountable for recommendation governance, disclosures, professional boundaries, and incident review.")
    add_bullet(doc, "Domain advisors: supply reviewed knowledge where advice could create safety, legal, insurance, financial, or construction risk.")

    heading(doc, "27. Major risks and mitigations", 2)
    add_table(doc, ["Risk", "Failure mode", "Mitigation"], [
        ("Vision breadth", "Too many journeys and utilities dilute quality", "Use the wedge, roadmap phases, and not-now list"),
        ("Cold start", "Users abandon before the product knows enough", "Import artifacts; show value early; progressive enrichment"),
        ("Weak data", "Confident guidance rests on stale or inferred facts", "Provenance, confidence, confirmation, conservative fallback"),
        ("Alert fatigue", "Users ignore important messages", "Priority model, digests, controls, and relevance feedback"),
        ("AI overreach", "Fluent output exceeds evidence or authority", "Grounding, tiered controls, boundaries, and auditability"),
        ("Commercial conflict", "Provider or product revenue erodes trust", "Independent analysis and visible disclosure"),
        ("Workflow shallowness", "Major events become generic checklists", "Launch fewer journeys with real state and dependencies"),
        ("Marketplace quality", "Poor provider outcomes damage the platform", "Qualification, scope clarity, outcome capture, recourse"),
        ("Privacy / transfer", "Household data is overshared or transferred", "Role-based permission and property/household separation"),
    ], [1700, 3350, 4310], font_size=8.5)

    heading(doc, "28. Decisions and validation agenda", 2)
    add_table(doc, ["Decision", "Current recommendation", "Evidence needed"], [
        ("Initial segment", "Recent buyers / first 24 months", "Problem intensity, acquisition economics, activation, retention"),
        ("First artifact", "Inspection report", "Availability, extraction quality, user trust, predictive value"),
        ("First-value output", "Personalized 90-day plan", "Comprehension, novelty, action completion"),
        ("First decision tool", "Repair/replace or quote evaluation", "Demand frequency, data sufficiency, willingness to pay"),
        ("First major event", "Buying and Moving In", "Journey completion and recurring-loop conversion"),
        ("Second event candidate", "Preparing to Sell", "Reuse, partner channel, monetization, customer urgency"),
        ("Business model", "Subscription plus carefully governed outcome-based services", "Trust impact, unit economics, disclosure comprehension"),
        ("North-star definition", "Early identification and successful completion", "Reliable event taxonomy, importance and completion criteria"),
    ], [1900, 3300, 4160], font_size=8.7)
    add_callout(doc, "Leadership requirement", "Resolve the launch wedge, first-value experience, initial decision module, first event, and monetization guardrails before scaling surface area.", PALE_GOLD, GOLD)

    heading(doc, "Appendix A — Feature definition test", 1)
    add_section_intro(doc, "Every proposed or existing capability should answer the following before it is retained, built, or prominently displayed.")
    questions = [
        "Which customer job does it primarily serve, and which jobs does it support secondarily?",
        "What observable homeowner outcome does it produce?",
        "What action should the user take?",
        "What context makes ContractToCozy better than a generic alternative?",
        "What information does the interaction add back to the Living Home Record?",
        "How does the capability become more valuable over time?",
        "What is its recommendation safety tier and required governance?",
        "How will usefulness, completion, and failure be measured?",
        "What happens when data is missing, conflicting, or low-confidence?",
        "What existing surface or capability should it replace, simplify, or reuse?",
    ]
    for idx, q in enumerate(questions):
        add_number(doc, q, restart=(idx == 0))
    heading(doc, "Worked example: appliance record", 2)
    add_para(doc, "Weak definition: Allows users to store appliance details.", italic=True, color=MUTED)
    add_para(doc, "Strong definition: Helps ContractToCozy identify warranty deadlines, maintenance needs, recall exposure, replacement timing, and repair-versus-replace recommendations; preserves evidence for claims, moving, and sale preparation.")
    add_table(doc, ["Dimension", "Contribution"], [
        ("Primary job", "Stay Ahead through maintenance, warranty, and recall detection"),
        ("Secondary jobs", "Decide through repair/replace; Major Moments through claim, move, and sale support"),
        ("User action", "Confirm the appliance, resolve a deadline, maintain, repair, or plan replacement"),
        ("System learning", "Model, age, coverage, service history, costs, provider, and outcome"),
        ("Success", "Relevant action is taken early or a better-supported decision is recorded"),
    ], [2250, 7110])

    heading(doc, "Appendix B — Recommendation launch checklist", 1)
    for item in [
        "The recommendation has a named homeowner outcome and owner.",
        "Required property facts, provenance, freshness, and confidence are defined.",
        "Options and material tradeoffs are represented fairly.",
        "The rationale is understandable without specialist knowledge.",
        "Uncertainty and missing information are visible.",
        "The user can correct, defer, dismiss, modify, or escalate.",
        "Commercial relationships and ranking influences are disclosed.",
        "Professional-advice and emergency boundaries are implemented.",
        "Low-confidence and system-failure fallbacks are safe and useful.",
        "The user’s decision and subsequent outcome can be recorded.",
        "Quality, fairness, complaint, reversal, and outcome metrics are observable.",
        "Support and incident-response paths are ready.",
    ]:
        add_bullet(doc, "☐ " + item)

    heading(doc, "Appendix C — Major-event design template", 1)
    add_table(doc, ["Field", "Definition prompt"], [
        ("Event", "What significant homeowner situation begins the journey?"),
        ("Target outcome", "What does complete and successful mean?"),
        ("Entry signals", "How does ContractToCozy know the event has begun?"),
        ("Stages", "What stable phases help the homeowner understand progress?"),
        ("Milestones", "What observable achievements mark advancement?"),
        ("Dependencies", "Which steps, documents, decisions, or people unblock others?"),
        ("Decision points", "Where should decision intelligence be invoked?"),
        ("Adaptive logic", "How does the next step change with context or outcome?"),
        ("Documents", "What evidence is needed, produced, or preserved?"),
        ("Participants", "Who owns, approves, advises, performs, or receives each step?"),
        ("Exceptions", "What urgent, disputed, delayed, denied, or failed paths exist?"),
        ("Completion", "What closes the event and what moves into the Living Home Record?"),
        ("Success metrics", "How will progress, completion, trust, and economic value be measured?"),
    ], [2000, 7360], font_size=9.0)

    heading(doc, "Appendix D — Product language system", 1)
    add_table(doc, ["Use", "Preferred language", "Avoid"], [
        ("Job 1", "Stay Ahead", "Daily"),
        ("Job 2", "Decide With Confidence", "Tools or calculators as the primary promise"),
        ("Job 3", "Navigate Major Moments", "Life Events as a generic feature bucket"),
        ("Record", "Living Home Record / home history", "Digital filing cabinet"),
        ("Attention", "What needs your attention", "Alerts center"),
        ("Urgency", "Now, Soon, Plan, Consider", "Alarmist or unexplained severity"),
        ("AI", "Guidance grounded in your home", "AI-powered as the main customer benefit"),
        ("Recommendation", "Based on these facts…", "The system knows best"),
    ], [1700, 3250, 4410], font_size=9.0)

    heading(doc, "Appendix E — Glossary", 1)
    glossary = [
        ("Attention item", "A ranked, actionable signal that explains relevance, timing, outcome, confidence, and next step."),
        ("Customer job", "The progress a homeowner is trying to make when engaging the product."),
        ("Decision intelligence", "Property-specific comparison and recommendation using evidence, tradeoffs, preferences, and uncertainty."),
        ("Living Home Record", "The permissioned, longitudinal property context accumulated through documents, facts, work, decisions, and outcomes."),
        ("Major moment", "A significant, multi-step home event requiring coordinated progress over time."),
        ("Recommendation contract", "The required fields and controls that make product guidance useful, explainable, and governable."),
        ("Signal", "A detected change, deadline, risk, opportunity, eligibility condition, or incomplete obligation."),
        ("Successful action", "A completed, intentionally deferred, or deliberately dismissed item with an understood outcome and recorded state."),
    ]
    for term, definition in glossary:
        add_para(doc, f"{term}. {definition}", bold_lead=f"{term}.")

    heading(doc, "Closing product doctrine", 1)
    add_callout(doc, "The standard", "ContractToCozy should help the homeowner notice earlier, choose more confidently, complete more successfully, and preserve what the home has learned.", PALE_TEAL, TEAL)
    add_para(doc, "The product wins when every new fact improves a future action, every completed action strengthens the home’s history, and every recommendation earns rather than assumes the homeowner’s trust.")

    # Set document metadata and consistent headers/footers after all sections exist.
    props = doc.core_properties
    props.title = "ContractToCozy Product Framework"
    props.subject = "Comprehensive product strategy and operating model"
    props.author = "ContractToCozy"
    props.keywords = "product strategy, homeowner jobs, home intelligence, roadmap, metrics, trust"
    props.comments = "Version 1.0 — July 2026"
    set_headers_footers(doc)

    # Widows/orphans and compatibility-friendly save.
    for p in doc.paragraphs:
        p.paragraph_format.widow_control = True
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
